from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, Any, Optional, Union
from pydantic import BaseModel


class NoteDrafter(ABC):
    """
    Base class for generating procedure notes from extracted data.
    Supports markdown and docx.
    """
    
    def __init__(
        self,
        template_config_path: Optional[Union[str, Path]] = None,
        procedure_type: str = "procedure"
    ):
        """
        Args:
            template_config_path: Path to drafter YAML config (optional)
            procedure_type: Type of procedure (e.g., 'ercp', 'colonoscopy')
        """
        self.template_config_path = Path(template_config_path) if template_config_path else None
        self.procedure_type = procedure_type
    
    @abstractmethod
    def get_section_order(self) -> list[str]:
        """Return ordered list of section names for the final note.
        Returns:
            List of section names in the order they should appear
        """
        pass
    
    @abstractmethod
    def format_impressions(self, impressions: Any) -> list[str]:
        """Format impressions data into a list of strings."""
        pass
    
    @abstractmethod
    def format_recommendations(self, recommendations: Any, extracted_data: Dict[str, Any]) -> list[str]:
        """Format recommendations data into a list of strings."""
        pass
    
    def render_markdown(
        self,
        extracted_data: Union[BaseModel, Dict[str, Any]],
        rendered_sections: Optional[Dict[str, str]] = None
    ) -> str:
        """Render note as markdown text.
        Args:
            extracted_data: Extracted field data (Pydantic model or dict)
            rendered_sections: Pre-rendered sections from templates (optional)
        Returns:
            Markdown-formatted note text
        """
        # Convert Pydantic model to dict
        if isinstance(extracted_data, BaseModel):
            data_dict = extracted_data.model_dump()
        else:
            data_dict = extracted_data
        
        sections = rendered_sections or {}
        note_parts = []
        section_order = self.get_section_order()
        for section_name in section_order:
            if section_name == 'impressions':
                impressions_text = self._render_impressions_markdown(data_dict)
                if impressions_text:
                    note_parts.append(impressions_text)
            elif section_name == 'recommendations':
                recommendations_text = self._render_recommendations_markdown(data_dict)
                if recommendations_text:
                    note_parts.append(recommendations_text)
            else:
                # Use pre-rendered section
                if section_name in sections:
                    rendered_text = sections[section_name].strip()
                    if rendered_text:
                        header = section_name.replace('_', ' ').upper()
                        note_parts.append(f"## {header}\n{rendered_text}")
        
        return "\n\n".join(note_parts)
    
    def render_docx(
        self,
        extracted_data: Union[BaseModel, Dict[str, Any]],
        rendered_sections: Optional[Dict[str, str]] = None
    ):
        """
        Args:
            extracted_data: Extracted field data (Pydantic model or dict)
            rendered_sections: Pre-rendered sections from templates (optional)
        Returns:
            docx.Document
        """
        from docx import Document
        from docx.shared import Pt

        if isinstance(extracted_data, BaseModel):
            data_dict = extracted_data.model_dump()
        else:
            data_dict = extracted_data
        
        sections = rendered_sections or {}
        
        doc = Document()
        doc.add_heading(f'{self.procedure_type.upper()} Procedure Note', level=1)
        
        section_order = self.get_section_order()
        for section_name in section_order:
            if section_name == 'impressions':
                impressions_list = self.format_impressions(data_dict.get('impressions'))
                if impressions_list:
                    doc.add_heading('Impressions', level=2)
                    for i, item in enumerate(impressions_list, start=1):
                        p = doc.add_paragraph(f"{i}. {item}")
                        p.paragraph_format.space_after = Pt(0)
            elif section_name == 'recommendations':
                recommendations_list = self.format_recommendations(
                    data_dict.get('recommendations'),
                    data_dict
                )
                if recommendations_list:
                    doc.add_heading('Recommendations', level=2)
                    for i, item in enumerate(recommendations_list, start=1):
                        p = doc.add_paragraph(f"{i}. {item}")
                        p.paragraph_format.space_after = Pt(0)
            else:
                if section_name in sections:
                    rendered_text = sections[section_name].strip()
                    if rendered_text:
                        header = section_name.replace('_', ' ').title()
                        doc.add_heading(header, level=2)
                        doc.add_paragraph(rendered_text)
        return doc
    
    def _render_impressions_markdown(
        self,
        data_dict: Dict[str, Any],
    ) -> str:
        impressions_list = self.format_impressions(data_dict.get('impressions'))
        if impressions_list:
            lines = ["## IMPRESSIONS"]
            for i, item in enumerate(impressions_list, start=1):
                lines.append(f"{i}. {item}")
            return "\n".join(lines)
        return ""
    
    def _render_recommendations_markdown(
        self,
        data_dict: Dict[str, Any],
    ) -> str:
        recommendations_list = self.format_recommendations(
            data_dict.get('recommendations'),
            data_dict
        )
        if recommendations_list:
            lines = ["## RECOMMENDATIONS"]
            for i, item in enumerate(recommendations_list, start=1):
                lines.append(f"{i}. {item}")
            return "\n".join(lines)
        return ""
