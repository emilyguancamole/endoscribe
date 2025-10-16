import pandas as pd
from drafters.base import EndoscopyDrafter
from docx import Document
from docx.shared import Pt
import re

from drafters.utils import find_terms_spacy

class ColonoscopyDrafter(EndoscopyDrafter):
    def __init__(self, sample, pred_df, patients_df, procedures_df, polyp_df):
        super().__init__(sample, pred_df, patients_df, procedures_df, polyp_df)

        self.colon_sample_df = self.sample_df
        self.polyp_sample_df = (
            polyp_df.loc[[sample]] if sample in polyp_df.index else pd.DataFrame()
        )
        self.polyp_count = int(self.colon_sample_df.get('polyp_count', 0))
    
    def get_indications(self):
        age, sex, indication_str = super().get_indications()
        return f"{age} year old {sex} here for a Colonoscopy for {indication_str}."

    def construct_recall(self):
        ''' 
        Follows algo to construct recall recommendation for colonoscopy.
        '''
        count = int(self.colon_sample_df['polyp_count'])
        if count == 0 or self.polyp_sample_df.empty:
            return "Return in 10 years for colonoscopy."
        if all(self.polyp_sample_df['size_max_mm'] < 10):
            if 1<=count<=2:
                return "Return in 20 years if hyperplastic polyps on pathology.\nReturn in 7 years if tubular adenoma on pathology.\nReturn in 5 years if sessile serrated polyp (SSP) on pathology.\nReturn in 3 years if adenoma with villous or tubulovillous or high grade dysplasia, OR sessile serrated polyp with dysplasia, OR traditional serrated adenoma on pathology."
            elif 3<=count<=4:
                return "Return in 10 years if hyperplastic polyps on pathology.\nReturn in 3-5 years if tubular adenoma or sessile serrated polyp (SSP) on pathology.\nReturn in 3 years if adenoma with villous or tubulovillous or high grade dysplasia, OR sessile serrated polyp with dysplasia, OR traditional serrated adenoma on pathology."
            elif 5<=count<=10:
                return "Return in 10 years if hyperplastic polyps on pathology.\nReturn in 3 years if tubular adenoma or sessile serrated polyp (SSP) on pathology.\nReturn in 3 years if adenoma with villous or tubulovillous or high grade dysplasia, OR sessile serrated polyp with dysplasia, OR traditional serrated adenoma on pathology."
            elif 10<count<20:
                return "Return in 10 years if hyperplastic polyps on pathology.\nReturn in 1 year if all are tubular adenoma.\nReturn in 3 years if adenoma with villous or tubulovillous or high grade dysplasia, OR sessile serrated polyp with dysplasia, OR traditional serrated adenoma on pathology."
            else:
                return "Return in 3 years for colonoscopy."
        else: # if any polyp >= 10mm
            return "Return in 3 years for colonoscopy."


    def construct_recommendations(self):
        '''
        Recommendations section for colonoscopy, including default and custom recommendations.
        Uses simple rules; uses negspacy to identify negated terms.
        '''
        rec = []
        colon_row = self.colon_sample_df
        if int(colon_row['polyp_count']) > 0:
            rec.append("Await biopsy results.")
        rec.append("Continue surveillance.")
        rec.append("Advance diet as tolerated.")
        rec.append("Resume current medications.")
        
        terms = find_terms_spacy(colon_row['findings'].lower(), ["ulcers", "ulcer"])
        if terms:
            rec.append("Avoid non-steroidal anti-inflammatory drugs.")

        if int(colon_row.get('age', 0)) < 75:
            rec.append("Continue age-appropriate colorectal cancer surveillance.")
        else:
            rec.append("Routine surveillance not typically recommended; will discuss on case-by-case basis.")

        rec.append("Follow up with referring physician.")
        return rec

    def draft_doc(self):
        ''' Create and return the report '''

        if self.sample not in self.pred_df.index:
            print(f"Sample {self.sample} not found in col predictions")
            return
        
        print(f"Creating colonoscopy report for '{self.sample}'")
        # sample_colon = self.pred_df.loc[self.sample] # df for the sample

        doc = Document()
        doc.add_heading(f'Report {self.sample}', level=1)

        doc.add_heading('Indications', level=2)
        indications = self.get_indications()
        if not indications:
            indications = self.sample_df.get('indications', 'unknown').replace('\\n', '\n')
        doc.add_paragraph(indications)

        doc.add_heading('Description of Procedure', level=2)
        description_extent = self.construct_description_procedure()
        doc.add_paragraph(description_extent)

        doc.add_heading('Prep Quality', level=2)
        pred_prep = self.construct_prep_report(self.colon_sample_df)
        doc.add_paragraph(pred_prep)

        doc.add_heading('Findings', level=2)
        doc.add_paragraph(self.colon_sample_df['findings'].replace('\\n', '\n'))

        doc.add_heading('Impressions', level=2)
        #* Split only at commas outside of quotes to get each list item
        impressions_text = str(self.colon_sample_df['impressions']).strip("[]")
        matches = re.findall(r'''(['"])(.*?)(\1)''', impressions_text) # extract strings, ignoring commas inside quotes
        impressions = [match[1] for match in matches]
        for i, item in enumerate(impressions, start=1):
            p = doc.add_paragraph(f"{i}. {item}")
            # print('\n',item)
            p.paragraph_format.space_after = Pt(0)

        doc.add_heading('Recommendations', level=2)
        recs = self.construct_recommendations()
        for i, rec in enumerate(recs, start=1): 
            p = doc.add_paragraph(f"{i}. {rec}")
            p.paragraph_format.space_after = Pt(0)

        doc.add_heading('Repeat Exam', level=2)
        repeat_exam = self.construct_recall()
        doc.add_paragraph(repeat_exam)
        return doc
    

###### Colonoscopy-specific functions ######

    def construct_description_procedure(self):
        extent = self.colon_sample_df.get('extent', 'unknown')

        report = ("After the risks, benefits and alternatives of the procedure were thoroughly explained, informed consent was obtained and confirmed.  Immediately prior to the procedure, a time-out was performed to verify the correct patient, procedure and site. A digital exam revealed no abnormalities of the rectum. The colonoscope was introduced through the anus and advanced to the {extent}.")

        return report.format(extent=extent)

    def construct_prep_report(self, row):
        ''' Construct "Prep Quality" section of report '''

        BBPSSimple = row.get('bbps_simple', 'unknown')
        BBPSRight = row.get('bbps_right', 'unknown')
        BBPSTransverse = row.get('bbps_transverse', 'unknown')
        BBPSLeft = row.get('bbps_left', 'unknown')
        BBPSTotal = row.get('bbps_total', 'unknown')
        
        report = ("The overall prep quality was {BBPSSimple}. "
                "The Boston Bowel Prep Score was Boston Scale Right colon {BBPSRight}, "
                "Transverse colon {BBPSTransverse}, Left colon {BBPSLeft}. "
                "Total BBPS = {BBPSTotal}.")
        
        return report.format(BBPSSimple=BBPSSimple,
                            BBPSRight=BBPSRight, 
                            BBPSTransverse=BBPSTransverse, 
                            BBPSLeft=BBPSLeft, 
                            BBPSTotal=BBPSTotal)

