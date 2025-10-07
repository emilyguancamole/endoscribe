from drafters.base import EndoscopyDrafter
from docx import Document
from docx.shared import Pt
import re

from drafters.utils import find_terms_spacy

class EGDDrafter(EndoscopyDrafter):
    def construct_recommendations(self):
        rec = []
        sample_row = self.sample_df # todo add in when reextracted
        if sample_row.get('samples_taken', 'False') == 'True':
            rec.append("Follow up pathology results.")

        findings = self.sample_df.get('egd_findings', '').lower()
        # IF include any “gastritis”, “duodenitis” or “ulcers”, recommendation should be “avoid non-steroidal anti-inflammatory drugs”
        terms = find_terms_spacy(findings, ["gastritis", "duodenitis", "ulcer", "ulcers"])
        if terms:
            rec.append("Avoid non-steroidal anti-inflammatory drugs.")
        rec.extend(["Advance diet as tolerated.", "Resume current medications.", "Follow up with referring provider."])
        return rec
        
        # TODO 
        # - IF patient receive Barrett’s RFA or cryoablation, “follow standard Barrett’s ablation
        # post-operative therapy”
        # - IF patient received treatment for any bleeding, “continue PPI”
        # - IF patient received a PEG or PEG-J or PEJ tube, ***
        # o Run fluids
        # o Consult nutrition for tube feed initiation and education
        # Dilation
        # Stent placement

    def construct_recall(self):
        pass

    def draft_doc(self):
        # Find sample, if multiple, use the first one
        print(f"Creating EGD report for '{self.sample}'") 

        doc = Document()
        doc.add_heading(f'Report {self.sample}', level=1)

        doc.add_heading('Indications', level=2)
        indications = self.sample_df.get('indications', 'unknown').replace('\\n', '\n')
        doc.add_paragraph(indications)

        doc.add_heading('EGD Findings', level=2)
        doc.add_paragraph("A high-definition endoscope was advanced to the " + self.sample_df['extent'].replace('\\n', '\n'))
        # bold
        doc.add_paragraph('Esophagus:').runs[0].font.bold = True
        doc.add_paragraph(self.sample_df['esophagus'].replace('\\n', '\n'))
        doc.add_paragraph('Stomach:').runs[0].font.bold = True
        doc.add_paragraph(self.sample_df['stomach'].replace('\\n', '\n'))
        doc.add_paragraph('Duodenum:').runs[0].font.bold = True
        doc.add_paragraph(self.sample_df['duodenum'].replace('\\n', '\n'))
        doc.add_paragraph('EGD Findings').runs[0].font.bold = True
        doc.add_paragraph(self.sample_df['egd_findings'].replace('\\n', '\n'))

        doc.add_heading('Impressions', level=2)
        impressions_text = self.sample_df['impressions'].strip("[]") 
        matches = re.findall(r'''(['"])(.*?)(\1)''', impressions_text) # split at commas outside quotes to get each list item
        impressions = [match[1] for match in matches]
        for i, item in enumerate(impressions, start=1):
            p = doc.add_paragraph(f"{i}. {item}")
            p.paragraph_format.space_after = Pt(0)

        doc.add_heading('Recommendations', level=2)
        recs = self.construct_recommendations()
        for i, rec in enumerate(recs, start=1): 
            p = doc.add_paragraph(f"{i}. {rec}")
            p.paragraph_format.space_after = Pt(0)

        return doc