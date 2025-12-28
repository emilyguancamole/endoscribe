from drafters.base import EndoscopyDrafter
from docx import Document
from docx.shared import Pt
import re
from drafters.utils import add_bold_subheading
import os
import yaml
from templating.drafter_engine import build_report_sections


class ERCPDrafter(EndoscopyDrafter):
    PROCEDURE_VARIANT = "ercp_base"  # default; can be overridden later via arg or data

    def get_indications(self):
        age, sex, indication_str = super().get_indications()
        return f"{age} year old {sex} here for Endoscopic Retrograde Cholangiopancreatography (ERCP) for {indication_str}."

    def construct_recommendations(self):
        rec = []
        ercp_row = self.sample_df
        if ercp_row.get('samples_taken', False):
            rec.append("Follow up pathology results.")
        rec.append("Finish IV fluids now.")
        rec.append("Pain control as needed.")

        # Stent recommendations
        biliary_stent = ercp_row.get('biliary_stent_type', 'None')
        if biliary_stent != "None":
            if biliary_stent=="plastic biliary 5F":
                rec.append("Repeat ERCP in 6-8 weeks for stent removal/replacement.")
            elif biliary_stent=="plastic biliary 7F":
                rec.append("Repeat ERCP in 3-4 months for stent removal/replacement.")
            elif biliary_stent=="plastic biliary 10F":
                rec.append("Repeat ERCP in ~4-5 months for stent removal/replacement.")
            elif biliary_stent=="other biliary":
                rec.append("Repeat ERCP for stent removal/replacement as clinically indicated.")
            elif biliary_stent=="FCSEMS bengign" or biliary_stent=="FCSEMS unknown" and self.patients_df.loc[self.sample].get('relevant_co_morbidities_malignancy___cancer')=="False":
                rec.append("Repeat ERCP in 3 months for stent removal/replacement.")
            elif biliary_stent=="FCSEMS malignant" or biliary_stent=="FCSEMS unknown" and self.patients_df.loc[self.sample].get('relevant_co_morbidities_malignancy___cancer')=="True":
                rec.append("Repeat ERCP in 6 months for stent removal/replacement.")
            elif biliary_stent=="FCSEMS unknown":
                rec.append("Repeat ERCP in 6 months for stent removal/replacement.")
        if ercp_row.get('pd_stent'):
            rec.append("AXR in 2-4 weeks to confirm stent passage.")
        #todo If stent was placed and patient had indication of chronic pancreatitis/pancreatic duct strictures/pancreatic duct leaks: • Repeat ERCP for stent exchange in 3 months
            ### update when I understand if indications come from pre-written or speech

        rec.append("Follow up with referring provider.")
        return rec
    
    def construct_recall(self):
        pass

    def construct_impressions(self, impressions_raw):
        # If the loader already parsed impressions as a list, use it
        if isinstance(impressions_raw, (list, tuple)):
            impressions = [str(x).strip() for x in impressions_raw if str(x).strip()]
        else:
            try:
                import json
                parsed = json.loads(impressions_raw)
                if isinstance(parsed, list):
                    impressions = [str(x).strip() for x in parsed if str(x).strip()]
            except Exception: # Fall back to regex for quoted items and line splitting
                impressions_text = str(impressions_raw or '').strip()
                if impressions_text: # Try to extract quoted strings: "..." or '...'
                    matches = re.findall(r'''(['"])(.*?)(\1)''', impressions_text)
                    if matches:
                        impressions = [m[1].strip() for m in matches if m[1].strip()]
                    else:
                        # Split by newline and strip leading bullets or numbering
                        lines = [ln.strip() for ln in re.split(r"\r?\n", impressions_text) if ln.strip()]
                        cleaned = []
                        for ln in lines: # remove common bullets or leading numbering
                            ln2 = re.sub(r'^\s*[-*\u2022]\s*', '', ln)
                            ln2 = re.sub(r'^\s*\d+\.|^\s*\d+\)', '', ln2).strip()
                            if ln2: cleaned.append(ln2)
                        impressions = cleaned
        return impressions

    def draft_doc(self):
        # Find sample, if multiple, use the first one
        
        print(f"Creating ERCP report for '{self.sample}'")


        doc = Document()
        doc.add_heading(f'Report {self.sample}', level=1)

        # History and Description of Procedure from YAML template if available
        try:
            index_fp = os.path.normpath(os.path.join(os.path.dirname(__file__), "procedures/index.yaml"))
            with open(index_fp, "r") as f:
                registry = yaml.safe_load(f) or {}
            variant_key = getattr(self, "procedure_variant", self.PROCEDURE_VARIANT)
            cfg_rel_path = registry.get("ercp", {}).get(variant_key)
            if cfg_rel_path:
                cfg_fp = os.path.normpath(os.path.join(os.path.dirname(index_fp), cfg_rel_path))
                # Assemble data dict from available sources (LLM + RedCap CSVs)
                data = {}
                # LLM-extracted fields for ERCP
                data.update({k: v for k, v in self.sample_df.items()})
                # Patient and procedure meta if present (safe get by index)
                if self.sample in self.procedures_df.index:
                    data.update({k: self.procedures_df.loc[self.sample].get(k) for k in self.procedures_df.columns})
                if self.sample in self.patients_df.index:
                    data.update({k: self.patients_df.loc[self.sample].get(k) for k in self.patients_df.columns})

                sections = build_report_sections(cfg_fp, data)
                
                # Render all sections from YAML templates
                indications = sections.get("indications", "").strip()
                hist = sections.get("history", "").strip()
                desc = sections.get("description_of_procedure", "").strip()
                findings = sections.get("findings", "").strip()
                quality = sections.get("ercp_quality_metrics", "").strip()

                doc.add_heading('Indications', level=2)
                if indications:
                    doc.add_paragraph(indications)
                else:
                    indications_fallback = self.get_indications()
                    if not indications_fallback:
                        indications_fallback = self.sample_df.get('indications', 'unknown').replace('\\n', '\n')
                    doc.add_paragraph(indications_fallback)

                # 11/16 Save rendered sections to self for external packaging for the reviewer
                try:
                    self.rendered_sections = {
                        'indications': indications or indications_fallback,
                        'medications': sections.get('medications', '').strip(),
                        'monitoring': sections.get('monitoring', '').strip(),
                        'post_procedure_static': sections.get('post_procedure_static', '').strip(),
                        'history': hist,
                        'description_of_procedure': desc,
                        'findings': findings,
                        'ercp_quality_metrics': quality,
                    }
                except Exception:
                    self.rendered_sections = None

                # Static/non-LLM sections:
                post_procedure_static = sections.get("post_procedure_static", "").strip()

                if hist:
                    doc.add_heading('History', level=2)
                    doc.add_paragraph(hist)
                if desc:
                    doc.add_heading('Description of Procedure', level=2)
                    doc.add_paragraph(desc)
                if findings:
                    doc.add_heading('Findings', level=2)
                    paragraph = doc.add_paragraph()
                    add_bold_subheading(paragraph, findings)
                    if post_procedure_static:
                        doc.add_paragraph(post_procedure_static)
                if quality:
                    doc.add_heading('ERCP Quality Metrics', level=2)
                    doc.add_paragraph(quality)
        except Exception as e:
            print(f"[ERROR] YAML template render skipped: {e}")
            import traceback
            traceback.print_exc()
            self.handle_fallbacks(self, doc)

        doc.add_heading('Impressions', level=2)
        impressions_raw = self.sample_df.get('impressions', '')
        impressions = self.construct_impressions(impressions_raw)
        
        # Write numbered list
        for i, item in enumerate(impressions, start=1):
            p = doc.add_paragraph(f"{i}. {item}")
            p.paragraph_format.space_after = Pt(0)
        # store impressions section text
        try:
            self.rendered_sections = getattr(self, 'rendered_sections', {}) or {}
            self.rendered_sections['impressions'] = '\n'.join([f"{i}. {it}" for i, it in enumerate(impressions, start=1)])
        except Exception:
            pass
        
        doc.add_heading('Recommendations', level=2)
        recommendations = self.construct_recommendations()
        for i, item in enumerate(recommendations, start=1):
            p = doc.add_paragraph(f"{i}. {item}")
            p.paragraph_format.space_after = Pt(0)
        # store recommendations
        try:
            self.rendered_sections = getattr(self, 'rendered_sections', {}) or {}
            self.rendered_sections['recommendations'] = '\n'.join([f"{i}. {it}" for i, it in enumerate(recommendations, start=1)])
        except Exception:
            pass

        # Compose rendered_note as a single string ordered by major sections
        try:
            parts = []
            order = ['indications', 'medications', 'monitoring', 'pre_procedure_exam', 'history', 'description_of_procedure', 'findings', 'ercp_quality_metrics', 'impressions', 'recommendations']
            for k in order:
                v = (getattr(self, 'rendered_sections', {}) or {}).get(k)
                if v:
                    parts.append(f"## {k.replace('_', ' ').title()}\n" + v.strip())
            self.rendered_note = "\n\n".join(parts)
        except Exception:
            self.rendered_note = None

        # provenance
        try:
            self.provenance = {
                'template': cfg_fp if 'cfg_fp' in locals() else None,
                'template_source': os.path.normpath(os.path.join(os.path.dirname(__file__), 'procedures'))
            }
        except Exception:
            self.provenance = None

        return doc
    

    def handle_fallbacks(self, doc: Document) -> Document:
        # During config file errors, make non-fatal: proceed with legacy content (non-generated templates) if config missing or bad
        doc.add_heading('Indications', level=2)
        indications_fallback = self.get_indications()
        if not indications_fallback:
            indications_fallback = self.sample_df.get('chief_complaints', self.sample_df.get('indications', 'unknown')).replace('\\n', '\n')
        doc.add_paragraph(indications_fallback)

        # Fallback: Medications and Monitoring (use sample_df or static defaults)
        doc.add_heading('Medications', level=2)
        med_fallback = self.sample_df.get('current_medications', 'Refer to record of source.')
        if not med_fallback or str(med_fallback).strip() in ['', 'none', 'unknown']:
            med_fallback = 'Refer to record of source.'
        doc.add_paragraph(med_fallback)

        doc.add_heading('Monitoring', level=2)
        mon_fallback = self.sample_df.get('monitoring', 'Johns Hopkins Standard.')
        if not mon_fallback or str(mon_fallback).strip() in ['', 'none', 'unknown']:
            mon_fallback = 'Johns Hopkins Standard.'
        doc.add_paragraph(mon_fallback)

        # Legacy findings
        doc.add_heading('EGD Findings', level=2)
        paragraph = doc.add_paragraph()
        text = self.sample_df.get('egd_findings', 'N/A').replace('\\n', '\n')
        add_bold_subheading(paragraph, text)
        doc.add_heading('ERCP Findings', level=2)
        paragraph = doc.add_paragraph()
        text = self.sample_df.get('ercp_findings', 'N/A').replace('\\n', '\n')
        add_bold_subheading(paragraph, text)