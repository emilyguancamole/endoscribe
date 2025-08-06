from drafters.base import EndoscopyDrafter
from docx import Document
from docx.shared import Pt
import re
from drafters.utils import add_bold_text


class EUSDrafter(EndoscopyDrafter):
    def construct_recommendations(self):
        '''
        Recommendations section for EUS. Currently only default recommendations.
        '''
        rec = ["MRI/MRCP in 1 year", "EUS in 2 years"]
        return rec

    def construct_recall(self):
        ''' Default recall recommendation for now '''
        return "Return for an upper endoscopic ultrasound as clinically indicated."

    def draft_doc(self):
        ''' EUS report generation ''' # was def create_eus_report_docx(sample, pred_df):

        print(f"Creating EUS report for '{self.sample}'")     

        doc = Document()
        doc.add_heading(f'Report {self.sample}', level=1)

        doc.add_heading('Indications', level=2)
        indications = self.sample_df.get('indications', 'unknown').replace('\\n', '\n')
        doc.add_paragraph(indications)
        
        doc.add_heading('EUS Findings', level=2)
        paragraph = doc.add_paragraph()
        # print("SAMPLE DF\n", self.sample_df)
        text = self.sample_df['eus_findings'].replace('\\n', '\n')
        add_bold_text(paragraph, text)
        doc.add_heading('EGD Findings', level=2)
        paragraph = doc.add_paragraph()
        text = self.sample_df['egd_findings'].replace('\\n', '\n')
        add_bold_text(paragraph, text)

        doc.add_heading('Impressions', level=2)
        impressions_text = self.sample_df['impressions'].strip("[]") 
        matches = re.findall(r'''(['"])(.*?)(\1)''', impressions_text) # split at commas outside quotes to get each list item
        impressions = [match[1] for match in matches]
        for i, item in enumerate(impressions, start=1):
            p = doc.add_paragraph(f"{i}. {item}")
            p.paragraph_format.space_after = Pt(0)

        doc.add_heading('Recommendations', level=2)
        recs = self.construct_recommendations()
        for i, rec in enumerate(recs, start=1): 
            p = doc.add_paragraph(f"{i}. {rec}")
            p.paragraph_format.space_after = Pt(0)

        doc.add_heading('Repeat Exam', level=2)
        repeat_exam = self.construct_recall()
        doc.add_paragraph(repeat_exam)

        return doc
